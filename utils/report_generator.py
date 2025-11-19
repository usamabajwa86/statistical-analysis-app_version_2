from docx import Document

def generate_docx_report(summary, save_path):
    doc = Document()
    doc.add_heading("Statistical Analysis Report", 0)

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)

    doc.save(save_path)
    return save_path
